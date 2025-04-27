import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


input_data = list(map(str.strip, sys.stdin))
document = Document()

title = document.add_heading('Blood test', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

table = document.add_table(rows=len(input_data) + 1, cols=3)
table.style = 'Table Grid'

table.cell(0, 0).text = 'Indicator'
table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table.cell(0, 1).text = 'Norm'
table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table.cell(0, 2).text = 'Value'
table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for row_index, row_value in enumerate(input_data, start=1):
    table.cell(row_index, 0).text = row_value

set_bold(table.rows[0])

document.save('analysis.docx')
