from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, Mm
import sys

input_text = list(map(lambda line: line.replace("\n", ""), sys.stdin.readlines()))

doc = Document()

greeting = doc.add_paragraph(input_text[0])
greeting.style = doc.styles.add_style("GreetingStyle", WD_STYLE_TYPE.PARAGRAPH)
greeting_font = greeting.style.font
greeting_font.name = "Arial"
greeting_font.size = Pt(11)
greeting_font.italic = True
greeting.alignment = WD_ALIGN_PARAGRAPH.RIGHT

normal_style = doc.styles["Normal"]
normal_style.font.name = "Times New Roman"
normal_style.font.size = Pt(12)

for paragraph_text in input_text[1:-2]:
    paragraph = doc.add_paragraph(paragraph_text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Inches(0.5)

closing_line_1 = doc.add_paragraph(input_text[-2])
closing_line_1.alignment = WD_ALIGN_PARAGRAPH.LEFT
closing_format = closing_line_1.paragraph_format
closing_format.space_after = Mm(0)

closing_line_2 = doc.add_paragraph(input_text[-1])
closing_line_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
closing_line_2.style = doc.styles.add_style("ClosingBoldStyle", WD_STYLE_TYPE.PARAGRAPH)
closing_font = closing_line_2.style.font
closing_font.bold = True

doc.save("letter.docx")
