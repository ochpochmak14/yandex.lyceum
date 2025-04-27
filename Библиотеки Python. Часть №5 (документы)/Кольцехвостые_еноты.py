from docx import Document
import sys

doc = Document()

nested_structure = {}
for line in list(map(str.rstrip, sys.stdin)):
    indentation_level = len(line) - len(line.strip())

    if indentation_level == 0:
        main_heading = line
        nested_structure[main_heading] = {}
    elif indentation_level == 2:
        sub_heading = line.strip()
        nested_structure[main_heading][sub_heading] = {}
    elif indentation_level == 4:
        sub_sub_heading = line.strip()
        nested_structure[main_heading][sub_heading][sub_sub_heading] = []
    else:
        item = line.strip()
        nested_structure[main_heading][sub_heading][sub_sub_heading].append(item)

first_heading = True
for heading in nested_structure:
    if first_heading:
        doc.add_heading(heading, level=0)
        first_heading = False
    doc.add_paragraph(heading, style='List Bullet')

    for subheading in nested_structure[heading]:
        doc.add_paragraph(subheading, style='List Bullet 2')

        for subsubheading in nested_structure[heading][subheading]:
            doc.add_paragraph(subsubheading, style='List Bullet 3')

            for list_item in nested_structure[heading][subheading][subsubheading]:
                doc.add_paragraph(list_item, style='List Bullet 4')

doc.save("systematization.docx")
