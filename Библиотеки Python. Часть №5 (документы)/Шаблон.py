from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docxtpl import DocxTemplate


def template(name):
    document = Document()

    p = document.add_heading("Самостоятельная работа №{{work}}", 1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = document.add_heading("Вариант №{{variant}}", 2)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3 = document.add_paragraph("Задание 1\n{{task1}}")
    p3.style = "List Number"
    p4 = document.add_paragraph("Задание 2\n{{task2}}")
    p4.style = "List Number"
    p5 = document.add_paragraph("Задание 3\n{{task3}}")
    p5.style = "List Number"

    p6 = document.add_paragraph()
    style = p6.add_run("Выполнил: {{name}}")
    p6f = p6.paragraph_format
    p6f.first_line_indent = Inches(0.5)
    style.italic = True

    document.save(name)


def prepare_doc(nameOfTemolate, **kwargs):

    doc = DocxTemplate(nameOfTemolate)

    doc.render(kwargs)
    doc.save("res.docx")
